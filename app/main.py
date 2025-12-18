import sys
import os
import json
from pathlib import Path
from dataclasses import dataclass
from enum import Enum
from typing import Any, Set, Optional

import re
import unicodedata
import pandas as pd

from docx import Document
from docx.enum.section import WD_ORIENT

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QFileDialog,
    QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QListWidget, QTableView,
    QMessageBox, QComboBox, QLineEdit, QDialog,
    QFormLayout, QDialogButtonBox, QTabWidget,
    QAbstractItemView, QSplitter, QTextEdit,
)
from PySide6.QtCore import Qt, QAbstractTableModel, QModelIndex
from PySide6.QtGui import QPixmap, QTextCursor, QTextCharFormat, QColor


CONFIG_PATH = Path.home() / ".table_filter_engine.json"
STATE_PATH = Path.home() / ".table_filter_engine_state.pkl"
SERVICE_COLS = {"is_archived", "is_deleted"}


# ============================================================
#                 УТИЛИТЫ / ЗАГРУЗКА ТАБЛИЦ
# ============================================================

def resource_path(rel_path: str) -> Path:
    """Корректный путь до ресурсов и в dev, и в exe (PyInstaller)."""
    if hasattr(sys, "_MEIPASS"):
        base = Path(sys._MEIPASS)
    else:
        base = Path(__file__).resolve().parent
    return base / rel_path


def _load_from_docx_first_table(path: str) -> pd.DataFrame:
    """
    Берём первую "осмысленную" таблицу из docx.
    Важно: делаем прямоугольник — строки, где меньше колонок, добиваем пустыми.
    """
    doc = Document(path)
    table = None
    for t in doc.tables:
        if len(t.rows) >= 2 and len(t.rows[0].cells) >= 2:
            table = t
            break
    if table is None:
        raise ValueError("У документі Word не знайдено підходящої таблиці.")

    headers = [c.text.strip() for c in table.rows[0].cells]
    rows = []
    for r in table.rows[1:]:
        vals = [c.text.strip() for c in r.cells]
        if len(vals) < len(headers):
            vals += [""] * (len(headers) - len(vals))
        elif len(vals) > len(headers):
            # если где-то внезапно больше — расширяем хедеры
            extra = len(vals) - len(headers)
            headers += [f"Col {len(headers)+i+1}" for i in range(extra)]
        rows.append(vals)

    df = pd.DataFrame(rows, columns=headers)
    return df.fillna("")


def load_test_df(path: str) -> pd.DataFrame:
    """
    Универсальный загрузчик: csv/xlsx/xls/docx.
    Для csv: автоопределение sep.
    """
    p = Path(path)
    ext = p.suffix.lower()

    if ext == ".csv":
        df = pd.read_csv(path, dtype=str, sep=None, engine="python").fillna("")
        return df

    if ext in (".xlsx", ".xls"):
        df = pd.read_excel(path, dtype=str).fillna("")
        return df

    if ext == ".docx":
        return _load_from_docx_first_table(path)

    raise ValueError(f"Формат не підтримується: {ext}")


# ============================================================
#                        ФИЛЬТРЫ
# ============================================================

class Operator(str, Enum):
    CONTAINS = "contains"
    EQUALS = "equals"
    NOT_EQUALS = "not_equals"
    RANGE = "range"


@dataclass
class FilterCondition:
    column: str
    operator: Operator
    value: Any


def _to_datetime_series(series: pd.Series) -> pd.Series:
    # пытаемся вытащить dd.mm.yyyy из строк и распарсить
    s = series.astype(str)
    d = s.str.extract(r"(\d{2}\.\d{2}\.\d{4})")[0]
    return pd.to_datetime(d, format="%d.%m.%Y", errors="coerce")


def apply_filters(df: pd.DataFrame, conditions: list[FilterCondition]) -> pd.DataFrame:
    out = df
    for cond in conditions:
        if cond.column not in out.columns:
            continue
        ser = out[cond.column]

        if cond.operator == Operator.CONTAINS:
            v = str(cond.value)
            mask = ser.astype(str).str.contains(v, case=False, na=False)
            out = out[mask]

        elif cond.operator == Operator.EQUALS:
            v = cond.value
            # делаем мягкое сравнение строками
            mask = ser.astype(str).str.strip().eq(str(v).strip())
            out = out[mask]

        elif cond.operator == Operator.NOT_EQUALS:
            v = cond.value
            mask = ~ser.astype(str).str.strip().eq(str(v).strip())
            out = out[mask]

        elif cond.operator == Operator.RANGE:
            d_from, d_to = cond.value
            dser = _to_datetime_series(ser)
            mask = pd.Series(True, index=out.index)
            if d_from is not None:
                mask &= dser >= d_from
            if d_to is not None:
                mask &= dser <= d_to
            out = out[mask]

    return out


# ============================================================
#                     МОДЕЛЬ ДЛЯ QTableView
# ============================================================

class PandasTableModel(QAbstractTableModel):
    def __init__(
        self,
        df: pd.DataFrame,
        edit_callback=None,
        expiring_by5_indices: Optional[Set[Any]] = None,
        expired_indices: Optional[Set[Any]] = None,
        duplicate_indices: Optional[Set[Any]] = None,
        ors_warning_indices: Optional[Set[Any]] = None,
        ors_overdue_indices: Optional[Set[Any]] = None,
        col5_name: str | None = None,
        col7_name: str | None = None,
        col8_name: str | None = None,
    ):
        super().__init__()
        self.df = df if df is not None else pd.DataFrame()
        self.edit_callback = edit_callback

        self.expiring_by5_indices = expiring_by5_indices or set()
        self.expired_indices = expired_indices or set()
        self.duplicate_indices = duplicate_indices or set()
        self.ors_warning_indices = ors_warning_indices or set()
        self.ors_overdue_indices = ors_overdue_indices or set()

        self.col5_name = col5_name
        self.col7_name = col7_name
        self.col8_name = col8_name

    def update_df(
        self,
        df: pd.DataFrame,
        expiring_by5_indices: Optional[Set[Any]] = None,
        expired_indices: Optional[Set[Any]] = None,
        duplicate_indices: Optional[Set[Any]] = None,
        ors_warning_indices: Optional[Set[Any]] = None,
        ors_overdue_indices: Optional[Set[Any]] = None,
        col5_name: str | None = None,
        col7_name: str | None = None,
        col8_name: str | None = None,
    ):
        self.beginResetModel()
        self.df = df if df is not None else pd.DataFrame()
        self.expiring_by5_indices = expiring_by5_indices or set()
        self.expired_indices = expired_indices or set()
        self.duplicate_indices = duplicate_indices or set()
        self.ors_warning_indices = ors_warning_indices or set()
        self.ors_overdue_indices = ors_overdue_indices or set()
        self.col5_name = col5_name
        self.col7_name = col7_name
        self.col8_name = col8_name
        self.endResetModel()

    def rowCount(self, parent=QModelIndex()):
        return 0 if self.df is None else len(self.df.index)

    def columnCount(self, parent=QModelIndex()):
        return 0 if self.df is None else len(self.df.columns)

    def headerData(self, section, orientation, role=Qt.DisplayRole):
        if role != Qt.DisplayRole or self.df is None:
            return None
        if orientation == Qt.Horizontal:
            try:
                return str(self.df.columns[section])
            except Exception:
                return ""
        else:
            try:
                return str(self.df.index[section])
            except Exception:
                return ""

    def flags(self, index: QModelIndex):
        if not index.isValid():
            return Qt.NoItemFlags
        base = Qt.ItemIsSelectable | Qt.ItemIsEnabled
        if self.edit_callback is not None:
            # запретим редактировать service-колонки
            col_name = str(self.df.columns[index.column()])
            if col_name not in SERVICE_COLS:
                base |= Qt.ItemIsEditable
        return base

    def data(self, index: QModelIndex, role=Qt.DisplayRole):
        if not index.isValid() or self.df is None:
            return None

        r = index.row()
        c = index.column()

        try:
            orig_index = self.df.index[r]
            col_name = str(self.df.columns[c])
            val = self.df.iloc[r, c]
        except Exception:
            return None

        if role == Qt.DisplayRole:
            if pd.isna(val):
                return ""
            return str(val)

        # ----------- подсветка -----------
        if role == Qt.BackgroundRole:
            # Архив - зелёный весь ряд
            if "is_archived" in self.df.columns:
                try:
                    if bool(self.df.at[orig_index, "is_archived"]):
                        return QColor("#d6f5d6")
                except Exception:
                    pass

            # Удалённые - серый ряд (если вдруг показывают)
            if "is_deleted" in self.df.columns:
                try:
                    if bool(self.df.at[orig_index, "is_deleted"]):
                        return QColor("#eeeeee")
                except Exception:
                    pass

            # Просрочка по 5-й колонке - красный весь ряд
            if orig_index in self.expired_indices:
                return QColor("#ffb3b3")

            # Дубликат ПІБ - синий весь ряд
            if orig_index in self.duplicate_indices:
                return QColor("#cfe8ff")

            # Жёлтая клетка в 5-й колонке
            if self.col5_name and orig_index in self.expiring_by5_indices:
                if col_name == self.col5_name:
                    return QColor("#fff2a8")

            # ОРС warning/overdue: подсветка клеток в 7 и 8
            if self.col7_name and self.col8_name:
                if col_name in (self.col7_name, self.col8_name):
                    if orig_index in self.ors_overdue_indices:
                        return QColor("#ffb3b3")
                    if orig_index in self.ors_warning_indices:
                        return QColor("#fff2a8")

        return None

    def setData(self, index: QModelIndex, value, role=Qt.EditRole):
        if role != Qt.EditRole or not index.isValid() or self.df is None:
            return False

        r = index.row()
        c = index.column()
        try:
            orig_index = self.df.index[r]
            col_name = str(self.df.columns[c])
        except Exception:
            return False

        if col_name in SERVICE_COLS:
            return False

        new_val = "" if value is None else str(value)

        # обновим df текущий
        self.df.iat[r, c] = new_val

        # колбэк — синхронизировать с df_original в MainWindow
        if self.edit_callback:
            self.edit_callback(orig_index, col_name, new_val)

        self.dataChanged.emit(index, index, [Qt.DisplayRole, Qt.BackgroundRole])
        return True


# ============================================================
#                 ДІАЛОГ ДОДАВАННЯ РЯДКА
# ============================================================

class AddRowDialog(QDialog):
    def __init__(self, prosecutors: list[str] | None = None, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Додати новий запис")
        self.setModal(True)

        prosecutors = prosecutors or []
        layout = QFormLayout(self)

        self.prosecutor_cb = QComboBox(self)
        self.prosecutor_cb.addItem("")
        for p in sorted(prosecutors):
            self.prosecutor_cb.addItem(str(p))
        layout.addRow("Прокуратура:", self.prosecutor_cb)

        self.case_edit = QLineEdit(self)
        self.case_edit.setPlaceholderText("№ провадження, дата, кваліфікація, орган…")
        layout.addRow("№ провадження / кваліфікація:", self.case_edit)

        self.fabula_edit = QLineEdit(self)
        self.fabula_edit.setPlaceholderText("Коротка фабула…")
        layout.addRow("Фабула:", self.fabula_edit)

        self.pib_edit = QLineEdit(self)
        self.pib_edit.setPlaceholderText("Прізвище Ім'я По батькові")
        layout.addRow("ПІБ підозрюваного:", self.pib_edit)

        self.dob_edit = QLineEdit(self)
        self.dob_edit.setPlaceholderText("дд.мм.рррр")
        layout.addRow("Дата народження:", self.dob_edit)

        self.notice_date_edit = QLineEdit(self)
        self.notice_date_edit.setPlaceholderText("дд.мм.рррр")
        layout.addRow("Дата повідомлення підозри:", self.notice_date_edit)

        self.measure_edit = QLineEdit(self)
        self.measure_edit.setPlaceholderText("Тримання під вартою / застава / ухвала …")
        layout.addRow("Запобіжний захід:", self.measure_edit)

        self.stop_edit = QLineEdit(self)
        self.stop_edit.setPlaceholderText("Підстава, дата зупинення…")
        layout.addRow("Зупинення розслідування:", self.stop_edit)

        self.order_edit = QLineEdit(self)
        self.order_edit.setPlaceholderText("Дата, вих. №, слідчий, адресат…")
        layout.addRow("Доручення / клопотання:", self.order_edit)

        self.ors_edit = QLineEdit(self)
        self.ors_edit.setPlaceholderText("№ ОРС, дата заведення, категорія, орган…")
        layout.addRow("№ ОРС:", self.ors_edit)

        self.border_edit = QLineEdit(self)
        self.border_edit.setPlaceholderText("Так/Ні, дата отримання інформації…")
        layout.addRow("Перетин кордону:", self.border_edit)

        self.admin_edit = QLineEdit(self)
        self.admin_edit.setPlaceholderText("Так/Ні, стаття, дата…")
        layout.addRow("Адмін. відповідальність:", self.admin_edit)

        self.interpol_edit = QLineEdit(self)
        self.interpol_edit.setPlaceholderText("Дата оголошення, № картки Інтерполу…")
        layout.addRow("Міжнародний розшук:", self.interpol_edit)

        btn_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, self)
        btn_box.accepted.connect(self.accept)
        btn_box.rejected.connect(self.reject)
        layout.addRow(btn_box)

    def get_data(self) -> dict[str, str]:
        return {
            "prosecutor": self.prosecutor_cb.currentText().strip(),
            "case_info": self.case_edit.text().strip(),
            "fabula": self.fabula_edit.text().strip(),
            "pib": self.pib_edit.text().strip(),
            "dob": self.dob_edit.text().strip(),
            "notice_date": self.notice_date_edit.text().strip(),
            "measure": self.measure_edit.text().strip(),
            "stop_info": self.stop_edit.text().strip(),
            "order_info": self.order_edit.text().strip(),
            "ors_info": self.ors_edit.text().strip(),
            "border_info": self.border_edit.text().strip(),
            "admin_info": self.admin_edit.text().strip(),
            "interpol_info": self.interpol_edit.text().strip(),
        }


# ============================================================
#        ДІАЛОГ АНАЛІЗУ ЗБІГІВ З ІНШИМ ДОКУМЕНТОМ
# ============================================================

class MatchAnalysisDialog(QDialog):
    def __init__(self, parent=None, current_df: pd.DataFrame | None = None):
        super().__init__(parent)
        self.setWindowTitle("Аналіз збігів")
        self.resize(1400, 820)

        self.current_df = current_df
        self.left_df: pd.DataFrame | None = None
        self.right_text: str = ""
        self.right_df: pd.DataFrame | None = None

        self.pib_matches: list[tuple[int, str]] = []
        self.pib_unique_rows: pd.DataFrame | None = None

        self.ors_matches: list[tuple[int, str]] = []
        self.ors_unique_rows: pd.DataFrame | None = None

        self._unique_pos_index: dict[tuple[str, str], int] = {}
        self.current_mode: str = "pib"

        top = QHBoxLayout()
        self.btn_use_current = QPushButton("Використати поточну таблицю")
        self.btn_load_table = QPushButton("Завантажити таблицю…")
        self.btn_load_doc = QPushButton("Завантажити документ справа…")
        self.btn_find_matches = QPushButton("Знайти збіги")
        self.btn_find_matches.setEnabled(False)

        top.addWidget(self.btn_use_current)
        top.addWidget(self.btn_load_table)
        top.addWidget(self.btn_load_doc)
        top.addStretch()
        top.addWidget(self.btn_find_matches)

        self.btn_use_current.clicked.connect(self.use_current_table)
        self.btn_load_table.clicked.connect(self.load_table_left)
        self.btn_load_doc.clicked.connect(self.load_right_document)
        self.btn_find_matches.clicked.connect(self.find_matches)

        self.left_table = QTableView()
        self.left_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.left_table.setSelectionMode(QAbstractItemView.SingleSelection)
        self.left_table.horizontalHeader().setStretchLastSection(True)

        self.right_tabs = QTabWidget()
        self.right_text_edit = QTextEdit()
        self.right_text_edit.setReadOnly(True)

        self.right_table = QTableView()
        self.right_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.right_table.horizontalHeader().setStretchLastSection(True)

        self.right_tabs.addTab(self.right_text_edit, "Текст")
        self.right_tabs.addTab(self.right_table, "Таблиця")
        self.right_tabs.setTabEnabled(1, False)

        center_splitter = QSplitter(Qt.Horizontal)
        left_panel = QWidget()
        lp = QVBoxLayout(left_panel)
        lp.setContentsMargins(0, 0, 0, 0)
        lp.addWidget(self.left_table)

        right_panel = QWidget()
        rp = QVBoxLayout(right_panel)
        rp.setContentsMargins(0, 0, 0, 0)
        rp.addWidget(self.right_tabs)

        center_splitter.addWidget(left_panel)
        center_splitter.addWidget(right_panel)
        center_splitter.setStretchFactor(0, 3)
        center_splitter.setStretchFactor(1, 2)

        self.bottom_tabs = QTabWidget()
        self.bottom_tabs.currentChanged.connect(self.on_bottom_tab_changed)

        # --- ПІБ вкладка ---
        pib_tab = QWidget()
        pib_layout = QVBoxLayout(pib_tab)
        pib_layout.setContentsMargins(0, 0, 0, 0)
        bottom_pib_splitter = QSplitter(Qt.Horizontal)

        match_panel_pib = QWidget()
        mp_pib = QVBoxLayout(match_panel_pib)
        mp_pib.setContentsMargins(0, 0, 0, 0)
        mp_pib.addWidget(QLabel("Збіги ПІБ:"))
        self.list_matches_pib = QListWidget()
        mp_pib.addWidget(self.list_matches_pib)

        self.btn_export_matches_pib = QPushButton("Експорт збігів у Word/Excel/CSV")
        self.btn_export_matches_pib.setEnabled(False)
        mp_pib.addWidget(self.btn_export_matches_pib)

        unique_panel_pib = QWidget()
        up_pib = QVBoxLayout(unique_panel_pib)
        up_pib.setContentsMargins(0, 0, 0, 0)
        up_pib.addWidget(QLabel("Рядки, яких немає в документі (ПІБ):"))
        self.list_unique_pib = QListWidget()
        up_pib.addWidget(self.list_unique_pib)

        self.btn_export_unique_pib = QPushButton("Експорт унікальних у Word/Excel/CSV")
        self.btn_export_unique_pib.setEnabled(False)
        up_pib.addWidget(self.btn_export_unique_pib)

        bottom_pib_splitter.addWidget(match_panel_pib)
        bottom_pib_splitter.addWidget(unique_panel_pib)
        bottom_pib_splitter.setStretchFactor(0, 1)
        bottom_pib_splitter.setStretchFactor(1, 1)
        pib_layout.addWidget(bottom_pib_splitter)

        # --- ОРС вкладка ---
        ors_tab = QWidget()
        ors_layout = QVBoxLayout(ors_tab)
        ors_layout.setContentsMargins(0, 0, 0, 0)
        bottom_ors_splitter = QSplitter(Qt.Horizontal)

        match_panel_ors = QWidget()
        mp_ors = QVBoxLayout(match_panel_ors)
        mp_ors.setContentsMargins(0, 0, 0, 0)
        mp_ors.addWidget(QLabel("Збіги ОРС:"))
        self.list_matches_ors = QListWidget()
        mp_ors.addWidget(self.list_matches_ors)

        self.btn_export_matches_ors = QPushButton("Експорт збігів у Word/Excel/CSV")
        self.btn_export_matches_ors.setEnabled(False)
        mp_ors.addWidget(self.btn_export_matches_ors)

        unique_panel_ors = QWidget()
        up_ors = QVBoxLayout(unique_panel_ors)
        up_ors.setContentsMargins(0, 0, 0, 0)
        up_ors.addWidget(QLabel("Рядки, яких немає в документі (ОРС):"))
        self.list_unique_ors = QListWidget()
        up_ors.addWidget(self.list_unique_ors)

        self.btn_export_unique_ors = QPushButton("Експорт унікальних у Word/Excel/CSV")
        self.btn_export_unique_ors.setEnabled(False)
        up_ors.addWidget(self.btn_export_unique_ors)

        bottom_ors_splitter.addWidget(match_panel_ors)
        bottom_ors_splitter.addWidget(unique_panel_ors)
        bottom_ors_splitter.setStretchFactor(0, 1)
        bottom_ors_splitter.setStretchFactor(1, 1)
        ors_layout.addWidget(bottom_ors_splitter)

        self.bottom_tabs.addTab(pib_tab, "ПІБ")
        self.bottom_tabs.addTab(ors_tab, "ОРС")

        self.list_matches_pib.itemSelectionChanged.connect(lambda: self.on_match_selected("pib"))
        self.list_unique_pib.itemSelectionChanged.connect(lambda: self.on_unique_selected("pib"))
        self.btn_export_unique_pib.clicked.connect(lambda: self.export_unique_rows("pib"))
        self.btn_export_matches_pib.clicked.connect(lambda: self.export_matches_rows("pib"))

        self.list_matches_ors.itemSelectionChanged.connect(lambda: self.on_match_selected("ors"))
        self.list_unique_ors.itemSelectionChanged.connect(lambda: self.on_unique_selected("ors"))
        self.btn_export_unique_ors.clicked.connect(lambda: self.export_unique_rows("ors"))
        self.btn_export_matches_ors.clicked.connect(lambda: self.export_matches_rows("ors"))

        layout = QVBoxLayout(self)
        layout.addLayout(top)
        layout.addWidget(center_splitter, 3)
        layout.addWidget(self.bottom_tabs, 2)

        if self.current_df is not None:
            self.set_left_df(self.current_df)

    # -------------------- ЛЕВАЯ ТАБЛИЦА --------------------

    def set_left_df(self, df: pd.DataFrame):
        self.left_df = df.copy()
        model = PandasTableModel(self.left_df, edit_callback=None)
        self.left_table.setModel(model)
        self.left_table.horizontalHeader().setStretchLastSection(True)

        self.btn_find_matches.setEnabled(self.left_df is not None and len(self.right_text.strip()) > 0)

    def use_current_table(self):
        if self.current_df is None:
            QMessageBox.warning(self, "Помилка", "У головному вікні немає таблиці.")
            return
        self.set_left_df(self.current_df)

    def load_table_left(self):
        path, _ = QFileDialog.getOpenFileName(self, "Вибрати таблицю", "", "Таблиці (*.csv *.xlsx *.xls *.docx)")
        if not path:
            return
        try:
            df = load_test_df(path)
            self.set_left_df(df)
        except Exception as e:
            QMessageBox.critical(self, "Помилка", str(e))

    # -------------------- ВСПОМОГАТЕЛЬНЫЕ --------------------

    def _get_pib_series(self) -> pd.Series | None:
        if self.left_df is None:
            return None
        df = self.left_df

        pib_col = next((c for c in df.columns if "ПІБ" in str(c)), None)
        if pib_col is not None:
            return df[pib_col].astype(str)

        def find_col(substrings: list[str]) -> str | None:
            for col in df.columns:
                name = str(col)
                if any(sub in name for sub in substrings):
                    return col
            return None

        col_last = find_col(["Прізвище", "Прiзвище", "Прізвище (укр)"])
        col_first = find_col(["Ім'", "Ім’я", "Імя", "Имя"])
        col_patron = find_col(["По батьков", "По-батьков", "По батькові", "По-батькові"])

        if not any([col_last, col_first, col_patron]):
            return None

        parts: list[pd.Series] = []
        if col_last:
            parts.append(df[col_last].astype(str).str.strip())
        if col_first:
            parts.append(df[col_first].astype(str).str.strip())
        if col_patron:
            parts.append(df[col_patron].astype(str).str.strip())

        result = parts[0]
        for ser in parts[1:]:
            result = result + " " + ser

        return result.str.replace(r"\s+", " ", regex=True).str.strip()

    def _get_ors_series(self) -> pd.Series | None:
        if self.left_df is None:
            return None
        df = self.left_df

        def find_ors_col() -> str | None:
            candidates = ["№ ОРС", "№ОРС", "№ОРД", "ОРД. РД", "ОРД РД", "ОРД/РД", "ОРС №"]
            for col in df.columns:
                name = str(col)
                if any(sub in name for sub in candidates):
                    return col
            return None

        ors_col = find_ors_col()
        if ors_col:
            base = df[ors_col].astype(str)
        else:
            base = df.astype(str).agg(" ".join, axis=1)

        extracted = base.str.extract(r"(\d{5,10})", expand=False).fillna("")
        if extracted.eq("").all():
            return None
        return extracted

    # -------------------- ПРАВЫЙ ДОКУМЕНТ --------------------

    def load_right_document(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Вибрати документ", "",
            "Документи (*.docx *.txt *.csv *.xlsx);;Усі файли (*)"
        )
        if not path:
            return

        try:
            ext = Path(path).suffix.lower()
            text = ""
            table = None

            if ext == ".txt":
                text = Path(path).read_text(encoding="utf-8", errors="ignore")

            elif ext in (".csv", ".xlsx"):
                if ext == ".csv":
                    df = pd.read_csv(path, dtype=str, sep=None, engine="python").fillna("")
                else:
                    df = pd.read_excel(path, dtype=str).fillna("")
                table = df
                rows = df.apply(lambda r: " ".join(r.values.astype(str)), axis=1)
                text = "\n".join(rows)

            elif ext == ".docx":
                doc = Document(path)
                parts = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)

                rows_all = []
                for t in doc.tables:
                    for r in t.rows:
                        cells = [" ".join(p.text for p in c.paragraphs).strip() for c in r.cells]
                        rows_all.append(cells)
                        row_text = " ".join(cells).strip()
                        if row_text:
                            parts.append(row_text)

                if rows_all:
                    maxc = max(len(r) for r in rows_all)
                    norm = [r + [""] * (maxc - len(r)) for r in rows_all]
                    table = pd.DataFrame(norm, columns=[f"Col {i+1}" for i in range(maxc)])

                text = "\n".join(parts)

            else:
                raise ValueError("Формат не підтримується")

            self.right_text = text or ""
            self.right_text_edit.setPlainText(self.right_text)

            self.right_df = table
            if table is not None:
                model = PandasTableModel(table, edit_callback=None)
                self.right_table.setModel(model)
                self.right_tabs.setTabEnabled(1, True)
            else:
                self.right_table.setModel(None)
                self.right_tabs.setTabEnabled(1, False)

            self.btn_find_matches.setEnabled(self.left_df is not None and bool(self.right_text.strip()))

        except Exception as e:
            QMessageBox.critical(self, "Помилка", str(e))

    # -------------------- ВКЛАДКИ НИЗУ --------------------

    def on_bottom_tab_changed(self, index: int):
        self.current_mode = "pib" if index == 0 else "ors"

    # -------------------- НОРМАЛИЗАЦИЯ (ИСПРАВЛЕНО) --------------------

    def _extract_dob_safe(self, text: str) -> str:
        """Витягує дату у форматі дд.мм.рррр. Повертає '' якщо немає."""
        if text is None:
            return ""
        s = str(text).strip()
        if not s or s in ("-", "—"):
            return ""
        m = re.search(r"\d{2}\.\d{2}\.\d{4}", s)
        return m.group(0) if m else ""

    def _normalize_pib_flexible(self, value: str) -> str:
        """
        Гнучка нормалізація ПІБ:
        - прибирає апострофи/тире/пунктуацію
        - прибирає невидимі пробіли
        - вирівнює латиницю↔кирилицю (візуальні двійники)
        - casefold
        """
        if value is None:
            return ""
        s = str(value)

        s = unicodedata.normalize("NFKC", s)

        for sp in ("\u00A0", "\u200B", "\u202F", "\ufeff"):
            s = s.replace(sp, " ")

        s = s.translate(str.maketrans({
            "A": "А", "a": "а",
            "B": "В",
            "C": "С", "c": "с",
            "E": "Е", "e": "е",
            "H": "Н",
            "I": "І", "i": "і",
            "K": "К",
            "M": "М",
            "O": "О", "o": "о",
            "P": "Р", "p": "р",
            "T": "Т",
            "X": "Х", "x": "х",
            "Y": "У", "y": "у",
        }))

        s = re.sub(r"[’ʼ'`´\-–—−‐]", "", s)
        s = re.sub(r"[^\w\s]", " ", s, flags=re.UNICODE)
        s = re.sub(r"\s+", " ", s).strip()

        return s.casefold()

    def _normalize_text_for_search(self, text: str) -> str:
        return self._normalize_pib_flexible(text or "")

    # -------------------- ПОИСК --------------------

    def find_matches(self):
        if self.left_df is None or not self.right_text.strip():
            QMessageBox.warning(self, "Помилка", "Потрібна таблиця зліва та документ справа.")
            return

        self._unique_pos_index.clear()
        self._find_pib_matches()
        self._find_ors_matches()

        self.bottom_tabs.setCurrentIndex(0)
        self.current_mode = "pib"
        self.highlight_all_pib_matches()

        if not self.pib_matches and not self.ors_matches:
            QMessageBox.information(self, "Готово", "Збігів не знайдено.")

    def _find_pib_matches(self):
        pib_series = self._get_pib_series()

        self.pib_matches = []
        self.pib_unique_rows = None
        self.list_matches_pib.clear()
        self.list_unique_pib.clear()
        self.btn_export_unique_pib.setEnabled(False)
        self.btn_export_matches_pib.setEnabled(False)

        if pib_series is None or self.left_df is None:
            return

        dob_col = next((c for c in self.left_df.columns if "Дата народ" in str(c) or "Дата рожд" in str(c)), None)

        raw_lines = [ln.strip() for ln in self.right_text.splitlines() if ln.strip()]
        norm_lines = [self._normalize_text_for_search(ln) for ln in raw_lines]

        matched_idx: set[int] = set()

        for idx, raw_name in pib_series.items():
            name = str(raw_name).split(",", 1)[0].strip()
            if not name:
                continue

            norm_name = self._normalize_pib_flexible(name)
            if not norm_name:
                continue

            dob = ""
            if dob_col and dob_col in self.left_df.columns:
                dob = self._extract_dob_safe(self.left_df.at[idx, dob_col])

            if not dob:
                dob = self._extract_dob_safe(raw_name)

            found_count = 0

            if dob:
                for norm_ln, raw_ln in zip(norm_lines, raw_lines):
                    if norm_name in norm_ln and dob in raw_ln:
                        found_count += 1
            else:
                for norm_ln in norm_lines:
                    if norm_name in norm_ln:
                        found_count += norm_ln.count(norm_name)

            if found_count > 0:
                self.pib_matches.append((idx, name))
                if dob:
                    self.list_matches_pib.addItem(f"{idx}: {name} | ДН: {dob} ({found_count})")
                else:
                    self.list_matches_pib.addItem(f"{idx}: {name} ({found_count})")
                matched_idx.add(idx)

        self.pib_unique_rows = self.left_df[~self.left_df.index.isin(matched_idx)].copy()

        for idx, raw_name in pib_series.items():
            if idx not in matched_idx:
                self.list_unique_pib.addItem(f"{idx}: {raw_name}")

        self.btn_export_unique_pib.setEnabled(self.pib_unique_rows is not None and not self.pib_unique_rows.empty)
        self.btn_export_matches_pib.setEnabled(len(self.pib_matches) > 0)

    def _find_ors_matches(self):
        ors_series = self._get_ors_series()

        self.ors_matches = []
        self.ors_unique_rows = None
        self.list_matches_ors.clear()
        self.list_unique_ors.clear()
        self.btn_export_unique_ors.setEnabled(False)
        self.btn_export_matches_ors.setEnabled(False)

        if ors_series is None or self.left_df is None:
            return

        text_lower = self.right_text.lower()
        matched_idx: set[int] = set()

        for idx, raw_num in ors_series.items():
            num = str(raw_num).strip()
            if not num:
                continue

            num_lower = num.lower()
            if num_lower in text_lower:
                count = text_lower.count(num_lower)
                self.ors_matches.append((idx, num))
                self.list_matches_ors.addItem(f"{idx}: {num} ({count})")
                matched_idx.add(idx)

        self.ors_unique_rows = self.left_df[~self.left_df.index.isin(matched_idx)].copy()
        for idx, raw_num in ors_series.items():
            if idx not in matched_idx:
                self.list_unique_ors.addItem(f"{idx}: {raw_num}")

        self.btn_export_unique_ors.setEnabled(self.ors_unique_rows is not None and not self.ors_unique_rows.empty)
        self.btn_export_matches_ors.setEnabled(len(self.ors_matches) > 0)

    # -------------------- ПОДСВЕТКА --------------------

    def highlight_all_pib_matches(self):
        doc = self.right_text_edit.document()

        cursor = QTextCursor(doc)
        cursor.select(QTextCursor.Document)
        cursor.setCharFormat(QTextCharFormat())

        if not self.pib_matches or not self.right_text:
            return

        fmt_yellow = QTextCharFormat()
        fmt_yellow.setBackground(Qt.yellow)

        text_lower = self.right_text.lower()

        for _, name in self.pib_matches:
            name_lower = name.lower()
            start = 0
            while True:
                pos = text_lower.find(name_lower, start)
                if pos == -1:
                    break
                cursor = QTextCursor(doc)
                cursor.setPosition(pos)
                cursor.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, len(name))
                cursor.mergeCharFormat(fmt_yellow)
                start = pos + len(name)

    # -------------------- UI: выбор из списков --------------------

    def on_match_selected(self, mode: str):
        if mode == "pib":
            item = self.list_matches_pib.currentItem()
            matches = self.pib_matches
        else:
            item = self.list_matches_ors.currentItem()
            matches = self.ors_matches

        if not item or not matches:
            return

        idx_str, rest = item.text().split(":", 1)
        idx = int(idx_str)
        name = rest.strip()
        if name.endswith(")"):
            pos_brace = name.rfind("(")
            if pos_brace != -1:
                name = name[:pos_brace].strip()

        model = self.left_table.model()
        if model and self.left_df is not None:
            for r in range(model.rowCount()):
                if self.left_df.index[r] == idx:
                    qidx = model.index(r, 0)
                    self.left_table.scrollTo(qidx)
                    self.left_table.selectRow(r)
                    break

        self.scroll_to_in_text(name)

        if self.right_df is not None:
            self.highlight_in_right_table(name)

    def on_unique_selected(self, mode: str):
        if mode == "pib":
            item = self.list_unique_pib.currentItem()
        else:
            item = self.list_unique_ors.currentItem()

        if not item:
            return

        try:
            _, full_text = item.text().split(":", 1)
            full_text = full_text.strip()
        except ValueError:
            return

        if mode == "pib":
            name = full_text.split(",")[0].strip()
        else:
            m = re.search(r"\d+", full_text)
            name = m.group(0).strip() if m else full_text

        if not name:
            return

        name_lower = name.lower()
        text_lower = self.right_text.lower()

        positions = []
        start = 0
        while True:
            pos = text_lower.find(name_lower, start)
            if pos == -1:
                break
            positions.append(pos)
            start = pos + len(name)

        if not positions:
            QMessageBox.information(self, "Немає вхождень", f"У документі не знайдено:\n{name}")
            return

        key = (mode, name_lower)
        cur_idx = self._unique_pos_index.get(key, -1) + 1
        if cur_idx >= len(positions):
            cur_idx = 0
        self._unique_pos_index[key] = cur_idx
        pos = positions[cur_idx]

        doc = self.right_text_edit.document()
        cursor = QTextCursor(doc)
        cursor.select(QTextCursor.Document)
        cursor.setCharFormat(QTextCharFormat())

        fmt = QTextCharFormat()
        fmt.setBackground(QColor("#d9b3ff"))

        cursor = self.right_text_edit.textCursor()
        cursor.setPosition(pos)
        cursor.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, len(name))
        cursor.mergeCharFormat(fmt)

        self.right_text_edit.setTextCursor(cursor)
        self.right_text_edit.ensureCursorVisible()

    def scroll_to_in_text(self, name: str):
        if not self.right_text:
            return

        text_lower = self.right_text.lower()
        name_lower = name.lower()
        pos = text_lower.find(name_lower)
        if pos == -1:
            return

        self.highlight_all_pib_matches()

        cursor = self.right_text_edit.textCursor()
        cursor.setPosition(pos)
        cursor.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, len(name))

        fmt_sel = QTextCharFormat()
        fmt_sel.setBackground(Qt.lightGray)
        cursor.mergeCharFormat(fmt_sel)

        self.right_text_edit.setTextCursor(cursor)
        self.right_text_edit.ensureCursorVisible()

    def highlight_in_right_table(self, name: str):
        model = self.right_table.model()
        if model is None:
            return
        name_lower = name.lower()
        for r in range(model.rowCount()):
            for c in range(model.columnCount()):
                val = str(model.index(r, c).data())
                if name_lower in val.lower():
                    self.right_table.selectRow(r)
                    self.right_table.scrollTo(model.index(r, 0))
                    return

    # -------------------- экспорт --------------------

    def _format_df_for_export(self, df: pd.DataFrame) -> pd.DataFrame:
        out = df.copy()
        for c in SERVICE_COLS:
            if c in out.columns:
                out = out.drop(columns=[c])
        for col in out.columns:
            if pd.api.types.is_datetime64_any_dtype(out[col]):
                out[col] = out[col].dt.strftime("%d.%m.%Y").fillna("")
            elif pd.api.types.is_bool_dtype(out[col]):
                out[col] = out[col].map({True: "Так", False: "Ні"})
        return out

    def _export_df(self, df: pd.DataFrame, title: str):
        if df is None or df.empty:
            QMessageBox.information(self, "Немає даних", "Немає рядків для експорту.")
            return

        path, selected_filter = QFileDialog.getSaveFileName(
            self, title, "", "Word (*.docx);;Excel (*.xlsx);;CSV (*.csv)"
        )
        if not path:
            return

        try:
            df_out = self._format_df_for_export(df)

            if path.lower().endswith(".docx") or "Word" in selected_filter:
                doc = Document()
                section = doc.sections[0]
                section.orientation = WD_ORIENT.LANDSCAPE
                new_width, new_height = section.page_height, section.page_width
                section.page_width = new_width
                section.page_height = new_height

                table = doc.add_table(rows=1, cols=len(df_out.columns))
                table.style = "Table Grid"

                hdr_cells = table.rows[0].cells
                for j, col_name in enumerate(df_out.columns):
                    hdr_cells[j].text = str(col_name)

                for _, row in df_out.iterrows():
                    row_cells = table.add_row().cells
                    for j, col_name in enumerate(df_out.columns):
                        value = row[col_name]
                        row_cells[j].text = "" if pd.isna(value) else str(value)

                doc.save(path)

            elif path.lower().endswith(".xlsx") or "Excel" in selected_filter:
                df_out.to_excel(path, index=False)
            else:
                df_out.to_csv(path, index=False)

            QMessageBox.information(self, "OK", f"Файл збережено:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Помилка", str(e))

    def export_unique_rows(self, mode: str):
        unique_rows = self.pib_unique_rows if mode == "pib" else self.ors_unique_rows
        self._export_df(unique_rows, "Зберегти унікальні рядки")

    def export_matches_rows(self, mode: str):
        if self.left_df is None or self.left_df.empty:
            QMessageBox.information(self, "Немає даних", "Немає лівої таблиці.")
            return

        indices = [idx for idx, _ in (self.pib_matches if mode == "pib" else self.ors_matches)]
        if not indices:
            QMessageBox.information(self, "Немає даних", "Немає збігів для експорту.")
            return

        seen = set()
        ordered = []
        for i in indices:
            if i not in seen:
                seen.add(i)
                ordered.append(i)

        df_matches = self.left_df.loc[ordered].copy()
        self._export_df(df_matches, "Зберегти збіги")


# ============================================================
#                      ГОЛОВНЕ ВІКНО
# ============================================================

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Table Filter Engine")
        self.resize(1500, 900)

        self.df_original: pd.DataFrame | None = None
        self.df_current: pd.DataFrame | None = None

        self.conditions: list[FilterCondition] = []
        self.global_search_text: str = ""

        self.expired_indices: Set[Any] = set()
        self.expiring_by5_indices: Set[Any] = set()
        self.ors_warning_indices: Set[Any] = set()
        self.ors_overdue_indices: Set[Any] = set()
        self.duplicate_indices: Set[Any] = set()

        self.ors_warning_rows: Set[Any] = set()
        self.ors_overdue_rows: Set[Any] = set()

        self.col5_name: str | None = None
        self.col7_name: str | None = None
        self.col8_name: str | None = None

        self.view_mode: str = "main"
        self.current_file_path: str | None = None

        self.show_only_expiring = False  # чтобы не падало, даже если не используешь

        self._init_ui()
        self._load_last_state_or_file()

    def _init_ui(self):
        central = QWidget()
        root = QVBoxLayout(central)
        root.setContentsMargins(5, 5, 5, 5)

        top = QHBoxLayout()

        self.btn_load = QPushButton("📂 Відкрити")
        self.btn_load.clicked.connect(self.open_file)
        top.addWidget(self.btn_load)

        self.btn_add = QPushButton("➕ Додати рядок")
        self.btn_add.clicked.connect(self.add_row)
        self.btn_add.setEnabled(False)
        top.addWidget(self.btn_add)

        self.btn_export = QPushButton("💾 Експорт")
        self.btn_export.clicked.connect(self.export_file)
        self.btn_export.setEnabled(False)
        top.addWidget(self.btn_export)

        self.btn_match = QPushButton("🔍 Збіги / аналіз")
        self.btn_match.clicked.connect(self.open_match_dialog)
        self.btn_match.setEnabled(False)
        top.addWidget(self.btn_match)

        top.addStretch()

        top.addWidget(QLabel("Глобальний пошук:"))
        self.ed_search = QLineEdit()
        self.ed_search.setPlaceholderText("Пошук по всіх стовпцях…")
        self.ed_search.textChanged.connect(self.on_global_search)
        self.ed_search.setEnabled(False)
        top.addWidget(self.ed_search, stretch=2)

        self.tab_mode = QTabWidget()
        self.tab_mode.addTab(QWidget(), "Основні")
        self.tab_mode.addTab(QWidget(), "Архів")
        self.tab_mode.addTab(QWidget(), "Видалені")
        self.tab_mode.addTab(QWidget(), "Прострочені (строк)")
        self.tab_mode.addTab(QWidget(), "Не заведено ОРС (20 діб)")
        self.tab_mode.addTab(QWidget(), "Не заведено ОРС (прострочено)")
        self.tab_mode.currentChanged.connect(self.on_tab_changed)
        self.tab_mode.setTabPosition(QTabWidget.North)
        top.addWidget(self.tab_mode)

        root.addLayout(top)

        main_splitter = QSplitter(Qt.Horizontal)

        left = QVBoxLayout()
        left.setAlignment(Qt.AlignTop)

        lbl_p = QLabel("Фільтр по прокуратурі:")
        lbl_p.setStyleSheet("font-weight: bold;")
        left.addWidget(lbl_p)

        self.cb_prosecutor = QComboBox()
        self.cb_prosecutor.addItem("Усі прокуратури")
        self.cb_prosecutor.currentIndexChanged.connect(self.apply_all_filters)
        self.cb_prosecutor.setEnabled(False)
        left.addWidget(self.cb_prosecutor)

        left.addSpacing(15)

        lbl_c = QLabel("Фільтр по стовпцю:")
        lbl_c.setStyleSheet("font-weight: bold;")
        left.addWidget(lbl_c)

        self.cb_column = QComboBox()
        self.cb_column.setEnabled(False)
        self.cb_column.currentIndexChanged.connect(self.on_column_changed)
        left.addWidget(self.cb_column)

        self.cb_operator = QComboBox()
        self.cb_operator.addItems(["містить", "дорівнює", "не дорівнює"])
        self.cb_operator.setEnabled(False)
        left.addWidget(self.cb_operator)

        self.ed_value = QLineEdit()
        self.ed_value.setPlaceholderText("Значення для фільтра…")
        self.ed_value.setEnabled(False)
        left.addWidget(self.ed_value)

        self.cb_value_choices = QComboBox()
        self.cb_value_choices.setVisible(False)
        self.cb_value_choices.currentIndexChanged.connect(self.on_value_choice_selected)
        left.addWidget(self.cb_value_choices)

        self.ed_date_from = QLineEdit()
        self.ed_date_from.setVisible(False)
        left.addWidget(self.ed_date_from)

        self.ed_date_to = QLineEdit()
        self.ed_date_to.setVisible(False)
        left.addWidget(self.ed_date_to)

        self.btn_add_condition = QPushButton("Додати умову")
        self.btn_add_condition.clicked.connect(self.add_condition_from_ui)
        self.btn_add_condition.setEnabled(False)
        left.addWidget(self.btn_add_condition)

        left.addSpacing(10)

        lbl_curr = QLabel("Поточні умови:")
        left.addWidget(lbl_curr)

        self.list_conditions = QListWidget()
        left.addWidget(self.list_conditions)

        self.btn_remove_condition = QPushButton("🗑 Видалити обрану умову")
        self.btn_remove_condition.clicked.connect(self.remove_selected_condition)
        self.btn_remove_condition.setEnabled(False)
        left.addWidget(self.btn_remove_condition)

        self.btn_clear_conditions = QPushButton("❌ Очистити всі умови")
        self.btn_clear_conditions.clicked.connect(self.clear_conditions)
        self.btn_clear_conditions.setEnabled(False)
        left.addWidget(self.btn_clear_conditions)

        left.addSpacing(10)

        self.btn_check_duplicates = QPushButton("Перевірити дублікати (ПІБ)")
        self.btn_check_duplicates.setEnabled(False)
        self.btn_check_duplicates.clicked.connect(self.on_check_duplicates_clicked)
        left.addWidget(self.btn_check_duplicates)

        left.addSpacing(10)

        lbl_ops = QLabel("Операції з рядками (за виділенням):")
        lbl_ops.setStyleSheet("font-weight: bold;")
        left.addWidget(lbl_ops)

        self.btn_to_archive = QPushButton("В архів")
        self.btn_to_archive.clicked.connect(self.move_selected_to_archive)
        self.btn_to_archive.setEnabled(False)
        left.addWidget(self.btn_to_archive)

        self.btn_from_archive = QPushButton("З архіву")
        self.btn_from_archive.clicked.connect(self.move_selected_from_archive)
        self.btn_from_archive.setEnabled(False)
        left.addWidget(self.btn_from_archive)

        self.btn_delete_rows = QPushButton("Видалити")
        self.btn_delete_rows.clicked.connect(self.delete_selected_rows)
        self.btn_delete_rows.setEnabled(False)
        left.addWidget(self.btn_delete_rows)

        self.btn_restore_rows = QPushButton("Відновити")
        self.btn_restore_rows.clicked.connect(self.restore_selected_rows)
        self.btn_restore_rows.setEnabled(False)
        left.addWidget(self.btn_restore_rows)

        self.list_conditions.itemDoubleClicked.connect(lambda _: self.remove_selected_condition())

        left_widget = QWidget()
        left_widget.setLayout(left)
        left_widget.setMinimumWidth(260)
        left_widget.setMaximumWidth(380)

        self.table_view = QTableView()
        self.table_view.setAlternatingRowColors(True)
        self.table_view.horizontalHeader().setStretchLastSection(True)
        self.table_view.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table_view.setSelectionMode(QAbstractItemView.ExtendedSelection)

        main_splitter.addWidget(left_widget)
        main_splitter.addWidget(self.table_view)
        main_splitter.setStretchFactor(0, 0)
        main_splitter.setStretchFactor(1, 1)

        root.addWidget(main_splitter)

        bottom_bar = QHBoxLayout()
        bottom_bar.setContentsMargins(4, 2, 4, 2)
        bottom_bar.setSpacing(6)

        self.btn_show_legend = QPushButton("Легенда кольорів")
        self.btn_show_legend.setFlat(True)
        self.btn_show_legend.setCursor(Qt.PointingHandCursor)
        self.btn_show_legend.setStyleSheet(
            "QPushButton { border: none; color: #555; font-size: 11px; "
            "text-decoration: underline; padding: 0 4px; } "
            "QPushButton:hover { color: #111; }"
        )
        self.btn_show_legend.clicked.connect(self.show_colors_legend)
        bottom_bar.addWidget(self.btn_show_legend)

        bottom_bar.addStretch()

        logo_label = QLabel()
        logo_path = resource_path("assets/national.png")
        if logo_path.exists():
            pm = QPixmap(str(logo_path))
            if not pm.isNull():
                pm = pm.scaledToHeight(36, Qt.SmoothTransformation)
                logo_label.setPixmap(pm)
        bottom_bar.addWidget(logo_label)

        copyright_label = QLabel("© Відділ 09/6")
        copyright_label.setStyleSheet("color: #555; font-size: 10px;")
        bottom_bar.addWidget(copyright_label)

        root.addLayout(bottom_bar)

        self.setCentralWidget(central)

    def show_colors_legend(self):
        QMessageBox.information(
            self,
            "Легенда кольорів",
            "Прострочений строк запобіжного заходу (6 міс. від ВІД) – червоний весь рядок\n"
            "Строк запобіжного заходу спливає (≤10 діб до 6 міс.) – жовта клітинка у 5-й колонці\n"
            "Не заведено ОРС (до 20 діб від дати у 7-й колонці) – жовті клітинки у 7-й та 8-й колонках\n"
            "Не заведено ОРС (прострочено >20 діб) – червоні клітинки у 7-й та 8-й колонках\n"
            "Дублікат ПІБ – синій фон всього рядка\n"
            "Архів – зелений фон всього рядка\n"
        )

    # -------------------- helpers --------------------

    def _is_date_like_column(self, series: pd.Series) -> bool:
        if pd.api.types.is_datetime64_any_dtype(series):
            return True
        try:
            return series.astype(str).str.contains(r"\d{2}\.\d{2}\.\d{4}").any()
        except Exception:
            return False

    def _save_last_file(self, path: str):
        try:
            CONFIG_PATH.write_text(json.dumps({"last_file": path}, ensure_ascii=False), encoding="utf-8")
        except Exception:
            pass

    def _save_state(self):
        if self.df_original is None:
            return
        try:
            self.df_original.to_pickle(STATE_PATH)
        except Exception:
            pass

    def _load_last_state_or_file(self):
        if STATE_PATH.exists():
            try:
                df = pd.read_pickle(STATE_PATH)
                self.current_file_path = None
                self._setup_dataframe(df, show_message=False)
                return
            except Exception:
                pass

        if not CONFIG_PATH.exists():
            return
        try:
            data = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
            last_path = data.get("last_file")
            if last_path and os.path.exists(last_path):
                self.load_table_from_path(last_path, show_message=False)
        except Exception:
            pass

    def _setup_dataframe(self, df: pd.DataFrame, show_message: bool):
        self.df_original = df
        self.df_current = df.copy()

        self.recalc_expiring_and_expired(show_popup=show_message)
        self.recalc_duplicate_marks(show_popup=show_message)

        model = PandasTableModel(
            self.df_current,
            edit_callback=self.on_cell_edited,
            expiring_by5_indices=self.expiring_by5_indices,
            expired_indices=self.expired_indices,
            duplicate_indices=self.duplicate_indices,
            ors_warning_indices=self.ors_warning_indices,
            ors_overdue_indices=self.ors_overdue_indices,
            col5_name=self.col5_name,
            col7_name=self.col7_name,
            col8_name=self.col8_name,
        )
        self.table_view.setModel(model)
        self.hide_service_columns()

        self.cb_prosecutor.setEnabled(True)
        self.cb_prosecutor.blockSignals(True)
        self.cb_prosecutor.clear()
        self.cb_prosecutor.addItem("Усі прокуратури")
        if "Прокуратура" in df.columns:
            for p in sorted(df["Прокуратура"].dropna().unique()):
                self.cb_prosecutor.addItem(str(p))
        self.cb_prosecutor.setCurrentIndex(0)
        self.cb_prosecutor.blockSignals(False)

        self.cb_column.setEnabled(True)
        self.cb_column.clear()
        for col in df.columns:
            if col not in SERVICE_COLS:
                self.cb_column.addItem(col)

        self.cb_operator.setEnabled(True)
        self.ed_value.setEnabled(True)
        self.btn_add_condition.setEnabled(True)
        self.btn_clear_conditions.setEnabled(True)
        self.btn_remove_condition.setEnabled(True)

        self.btn_add.setEnabled(True)
        self.btn_export.setEnabled(True)
        self.btn_match.setEnabled(True)
        self.ed_search.setEnabled(True)
        self.btn_check_duplicates.setEnabled(True)

        self.conditions.clear()
        self.list_conditions.clear()
        self.global_search_text = ""
        self.ed_search.clear()

        self.view_mode = "main"
        self.tab_mode.setCurrentIndex(0)
        self.update_action_buttons_state()

        self.on_column_changed(self.cb_column.currentIndex())

        self._save_state()

    # -------------------- загрузка --------------------

    def open_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Вибрати файл реєстру", "",
            "Таблиці (*.csv *.xlsx *.xls *.docx);;Усі файли (*)"
        )
        if not path:
            return
        self.load_table_from_path(path, show_message=True)

    def load_table_from_path(self, path: str, show_message: bool = True):
        try:
            df = load_test_df(path)

            if "is_archived" not in df.columns:
                df["is_archived"] = False
            if "is_deleted" not in df.columns:
                df["is_deleted"] = False

            self.current_file_path = path
            self._save_last_file(path)

            self._setup_dataframe(df, show_message=show_message)

            if show_message:
                QMessageBox.information(self, "OK", f"Файл завантажено:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Помилка завантаження", str(e))

    def hide_service_columns(self):
        model = self.table_view.model()
        if not isinstance(model, PandasTableModel):
            return
        df = model.df
        for name in ("is_archived", "is_deleted"):
            if name in df.columns:
                idx = df.columns.get_loc(name)
                self.table_view.setColumnHidden(idx, True)

    # -------------------- пересчет сроков --------------------

    def recalc_expiring_and_expired(self, show_popup: bool = True):
        self.expired_indices = set()
        self.expiring_by5_indices = set()
        self.ors_warning_indices = set()
        self.ors_overdue_indices = set()
        self.ors_warning_rows = set()
        self.ors_overdue_rows = set()

        if self.df_original is None:
            return

        df = self.df_original
        today = pd.Timestamp.today().normalize()
        cutoff_5 = pd.Timestamp(2025, 9, 1)

        self.col5_name = next((c for c in df.columns if "Запобіжний захід" in str(c)), None)
        self.col7_name = next(
            (c for c in df.columns if "Дата та вихідний № доручення" in str(c) or ("доручення" in str(c) and "вихідний" in str(c))),
            None,
        )
        self.col8_name = next((c for c in df.columns if "№ ОРС" in str(c)), None)

        if self.col5_name:
            ser5 = df[self.col5_name].astype(str)
            first_dates_str = ser5.str.extract(r"(\d{2}\.\d{2}\.\d{4})")[0]
            dates5 = pd.to_datetime(first_dates_str, format="%d.%m.%Y", errors="coerce")
            expiry_dates = dates5 + pd.DateOffset(months=6)

            for idx in df.index:
                d_exp = expiry_dates.loc[idx]
                if pd.isna(d_exp):
                    continue
                if d_exp < cutoff_5:
                    continue
                days_left = (d_exp - today).days
                if days_left < 0:
                    self.expired_indices.add(idx)
                elif 0 <= days_left <= 10:
                    self.expiring_by5_indices.add(idx)

        if self.col7_name and self.col8_name:
            ser7 = df[self.col7_name].astype(str)
            ser8 = df[self.col8_name].astype(str)

            d7 = pd.to_datetime(ser7.str.extract(r"(\d{2}\.\d{2}\.\d{4})")[0], format="%d.%m.%Y", errors="coerce")
            d8 = pd.to_datetime(ser8.str.extract(r"(\d{2}\.\d{2}\.\d{4})")[0], format="%d.%m.%Y", errors="coerce")

            for idx in df.index:
                base_date = d7.loc[idx]
                ors_date = d8.loc[idx]
                if pd.isna(base_date):
                    continue
                if not pd.isna(ors_date):
                    continue

                days_passed = (today - base_date).days
                if 0 <= days_passed <= 20:
                    self.ors_warning_indices.add(idx)
                    self.ors_warning_rows.add(idx)
                elif days_passed > 20:
                    self.ors_overdue_indices.add(idx)
                    self.ors_overdue_rows.add(idx)

        if show_popup:
            parts = []
            if self.expired_indices:
                parts.append(f"Прострочені строки (5-та колонка): {len(self.expired_indices)}")
            if self.expiring_by5_indices:
                parts.append(f"Строки, що спливають (≤10 діб, 5-та колонка): {len(self.expiring_by5_indices)}")
            if self.ors_warning_rows:
                parts.append(f"Не заведено ОРС (до 20 діб): {len(self.ors_warning_rows)}")
            if self.ors_overdue_rows:
                parts.append(f"Не заведено ОРС (прострочено): {len(self.ors_overdue_rows)}")
            if parts:
                QMessageBox.warning(self, "Увага", "\n".join(parts))

    # -------------------- дубликаты --------------------

    def recalc_duplicate_marks(self, show_popup: bool = True):
        old_count = len(self.duplicate_indices)
        self.duplicate_indices = set()

        if self.df_original is None:
            return

        df = self.df_original
        if "is_deleted" in df.columns:
            df = df[df["is_deleted"] == False]

        if df.empty:
            return

        pib_col = next((c for c in df.columns if "ПІБ" in str(c)), None)
        if pib_col is None:
            return

        full_series = df[pib_col].astype(str)
        name_series = full_series.str.split(",", n=1).str[0].str.strip()
        valid = name_series != ""
        name_valid = name_series[valid]
        if name_valid.empty:
            return

        counts = name_valid.value_counts()
        dup_names = set(counts[counts > 1].index)
        if not dup_names:
            return

        mask_dups = name_series.isin(dup_names)
        idxs = df.index[mask_dups].tolist()
        self.duplicate_indices.update(idxs)

        if show_popup and len(self.duplicate_indices) > old_count:
            QMessageBox.warning(self, "Дублікати", f"Виявлено {len(self.duplicate_indices)} запис(ів)-дублікат(ів) (за ПІБ).")

    # -------------------- поиск/фильтры --------------------

    def on_global_search(self, text: str):
        self.global_search_text = text.strip()
        self.apply_all_filters()

    def on_column_changed(self, index: int):
        if self.df_original is None or index < 0:
            return

        column = self.cb_column.itemText(index)
        if not column:
            return

        series = self.df_original[column]
        is_date_like = self._is_date_like_column(series)

        self.cb_operator.setVisible(True)
        self.ed_value.setVisible(True)

        if is_date_like:
            self.ed_date_from.setVisible(True)
            self.ed_date_to.setVisible(True)
            self.ed_date_from.setPlaceholderText("з дд.мм.рррр (можна не заповнювати)")
            self.ed_date_to.setPlaceholderText("по дд.мм.рррр (можна не заповнювати)")
        else:
            self.ed_date_from.setVisible(False)
            self.ed_date_to.setVisible(False)

        self.ed_date_from.clear()
        self.ed_date_to.clear()

        uniques = series.dropna().unique()
        if len(uniques) <= 50 or column in ("Стаття_ККУ", "Категорія_розшуку"):
            self.cb_value_choices.setVisible(True)
            self.cb_value_choices.clear()
            self.cb_value_choices.addItem("— оберіть значення —")
            for val in sorted(map(str, uniques)):
                self.cb_value_choices.addItem(val)
        else:
            self.cb_value_choices.setVisible(False)

    def on_value_choice_selected(self, index: int):
        if index <= 0:
            return
        self.ed_value.setText(self.cb_value_choices.currentText())

    def add_condition_from_ui(self):
        if self.df_original is None:
            return

        column = self.cb_column.currentText()
        if not column:
            return

        series = self.df_original[column]
        is_date_like = self._is_date_like_column(series)

        if is_date_like:
            from_text = self.ed_date_from.text().strip()
            to_text = self.ed_date_to.text().strip()
            if from_text or to_text:
                def parse_date(txt: str):
                    if not txt:
                        return None
                    try:
                        return pd.to_datetime(txt, format="%d.%m.%Y", dayfirst=True)
                    except Exception:
                        QMessageBox.warning(self, "Невірний формат дати", "Використовуйте формат дд.мм.рррр (наприклад, 05.01.2025).")
                        raise

                try:
                    d_from = parse_date(from_text)
                    d_to = parse_date(to_text)
                except Exception:
                    return

                cond = FilterCondition(column=column, operator=Operator.RANGE, value=(d_from, d_to))
                self.conditions.append(cond)

                label_from = from_text or "…"
                label_to = to_text or "…"
                self.list_conditions.addItem(f"{column}: {label_from} — {label_to}")

                self.ed_date_from.clear()
                self.ed_date_to.clear()
                self.apply_all_filters()
                return

        op_text = self.cb_operator.currentText()
        raw_value = self.ed_value.text().strip()
        if not op_text or not raw_value:
            return

        if op_text == "містить":
            operator = Operator.CONTAINS
        elif op_text == "дорівнює":
            operator = Operator.EQUALS
        else:
            operator = Operator.NOT_EQUALS

        cond = FilterCondition(column=column, operator=operator, value=raw_value)
        self.conditions.append(cond)
        self.list_conditions.addItem(f"{column} {op_text} {raw_value}")

        self.ed_value.clear()
        self.apply_all_filters()

    def remove_selected_condition(self):
        idx = self.list_conditions.currentRow()
        if idx < 0 or idx >= len(self.conditions):
            return
        del self.conditions[idx]
        self.list_conditions.takeItem(idx)
        self.apply_all_filters()

    def clear_conditions(self):
        self.conditions.clear()
        self.list_conditions.clear()
        self.apply_all_filters()

    def on_tab_changed(self, index: int):
        if index == 0:
            self.view_mode = "main"
        elif index == 1:
            self.view_mode = "archive"
        elif index == 2:
            self.view_mode = "deleted"
        elif index == 3:
            self.view_mode = "expired"
        elif index == 4:
            self.view_mode = "ors_warning"
        else:
            self.view_mode = "ors_overdue"

        self.update_action_buttons_state()
        self.apply_all_filters()

    def update_action_buttons_state(self):
        if self.view_mode == "main":
            self.btn_to_archive.setEnabled(True)
            self.btn_from_archive.setEnabled(False)
            self.btn_delete_rows.setEnabled(True)
            self.btn_restore_rows.setEnabled(False)
        elif self.view_mode == "archive":
            self.btn_to_archive.setEnabled(False)
            self.btn_from_archive.setEnabled(True)
            self.btn_delete_rows.setEnabled(True)
            self.btn_restore_rows.setEnabled(False)
        elif self.view_mode == "deleted":
            self.btn_to_archive.setEnabled(False)
            self.btn_from_archive.setEnabled(False)
            self.btn_delete_rows.setEnabled(False)
            self.btn_restore_rows.setEnabled(True)
        else:
            self.btn_to_archive.setEnabled(True)
            self.btn_from_archive.setEnabled(True)
            self.btn_delete_rows.setEnabled(True)
            self.btn_restore_rows.setEnabled(True)

    def apply_all_filters(self):
        if self.df_original is None:
            return

        df = self.df_original.copy()

        if self.conditions:
            df = apply_filters(df, self.conditions)

        pros = self.cb_prosecutor.currentText()
        if pros and pros != "Усі прокуратури" and "Прокуратура" in df.columns:
            df = df[df["Прокуратура"] == pros]

        if self.global_search_text:
            text = self.global_search_text
            mask = df.apply(lambda col: col.astype(str).str.contains(text, case=False, na=False), axis=0).any(axis=1)
            df = df[mask]

        if "is_deleted" in df.columns:
            if self.view_mode == "main":
                df = df[df["is_deleted"] == False]
            elif self.view_mode == "archive":
                df = df[(df["is_deleted"] == False) & (df["is_archived"] == True)]
            elif self.view_mode == "deleted":
                df = df[df["is_deleted"] == True]
            elif self.view_mode == "expired":
                df = df[(df["is_deleted"] == False) & (df.index.isin(self.expired_indices))]
            elif self.view_mode == "ors_warning":
                df = df[(df["is_deleted"] == False) & (df.index.isin(self.ors_warning_rows))]
            elif self.view_mode == "ors_overdue":
                df = df[(df["is_deleted"] == False) & (df.index.isin(self.ors_overdue_rows))]

        self.df_current = df

        model = self.table_view.model()
        if isinstance(model, PandasTableModel):
            model.update_df(
                self.df_current,
                expiring_by5_indices=self.expiring_by5_indices,
                expired_indices=self.expired_indices,
                duplicate_indices=self.duplicate_indices,
                ors_warning_indices=self.ors_warning_indices,
                ors_overdue_indices=self.ors_overdue_indices,
                col5_name=self.col5_name,
                col7_name=self.col7_name,
                col8_name=self.col8_name,
            )
        else:
            self.table_view.setModel(PandasTableModel(self.df_current, edit_callback=self.on_cell_edited))

        self.hide_service_columns()

    # -------------------- синхронизация правок --------------------

    def on_cell_edited(self, orig_index, column_name: str, new_value):
        if self.df_original is None:
            return
        if orig_index in self.df_original.index and column_name in self.df_original.columns:
            self.df_original.at[orig_index, column_name] = new_value

        if column_name not in ("is_archived", "is_deleted"):
            self.recalc_expiring_and_expired(show_popup=False)
            self.recalc_duplicate_marks(show_popup=False)

        self._save_state()
        self.apply_all_filters()

    # -------------------- выделение --------------------

    def get_selected_indices(self) -> list[int]:
        if self.df_current is None:
            return []
        indices: set[int] = set()
        sel_model = self.table_view.selectionModel()
        if sel_model is not None:
            for idx in sel_model.selectedRows():
                try:
                    orig_index = self.df_current.index[idx.row()]
                    indices.add(orig_index)
                except Exception:
                    continue
        return list(indices)

    # -------------------- добавление --------------------

    def add_row(self):
        if self.df_original is None:
            return

        if "Прокуратура" in self.df_original.columns:
            prosecutors = sorted(self.df_original["Прокуратура"].dropna().unique())
        else:
            prosecutors = []

        dlg = AddRowDialog(prosecutors=prosecutors, parent=self)
        if dlg.exec() != QDialog.Accepted:
            return

        data = dlg.get_data()
        cols = list(self.df_original.columns)

        pib = data["pib"]
        dob = data["dob"]
        notice_date = data["notice_date"]
        pib_block = ", ".join([v for v in [pib, dob, notice_date] if v])

        new_id = None
        if "ID" in cols:
            try:
                max_id = pd.to_numeric(self.df_original["ID"], errors="coerce").max()
                if pd.isna(max_id):
                    max_id = 0
                new_id = int(max_id) + 1
            except Exception:
                new_id = len(self.df_original) + 1

        row: dict[str, object] = {}
        for col in cols:
            text_col = str(col)

            if col == "ID" and new_id is not None:
                row[col] = new_id
            elif text_col == "Прокуратура":
                row[col] = data["prosecutor"]
            elif "№ кримінального провадження" in text_col:
                row[col] = data["case_info"]
            elif text_col.strip() == "Фабула":
                row[col] = data["fabula"]
            elif "ПІБ підозрюваного" in text_col:
                row[col] = pib_block
            elif "Запобіжний захід" in text_col:
                row[col] = data["measure"]
            elif "Підстава, дата зупинення" in text_col:
                row[col] = data["stop_info"]
            elif "Дата та вихідний № доручення" in text_col:
                row[col] = data["order_info"]
            elif "№ ОРС, дата заведення" in text_col:
                row[col] = data["ors_info"]
            elif "Наявність інформації про перетин кордону" in text_col:
                row[col] = data["border_info"]
            elif "Притягнення до адмін" in text_col:
                row[col] = data["admin_info"]
            elif "Дата оголошення у міжнародний розшук" in text_col:
                row[col] = data["interpol_info"]
            elif col == "is_archived":
                row[col] = False
            elif col == "is_deleted":
                row[col] = False
            else:
                row[col] = ""

        new_row_df = pd.DataFrame([row], columns=self.df_original.columns)
        self.df_original = pd.concat([self.df_original, new_row_df], ignore_index=True)

        self.recalc_expiring_and_expired(show_popup=False)
        self.recalc_duplicate_marks(show_popup=True)
        self._save_state()
        self.apply_all_filters()

    # -------------------- операции с рядами --------------------

    def move_selected_to_archive(self):
        idxs = self.get_selected_indices()
        if not idxs:
            QMessageBox.information(self, "Архів", "Не вибрано жодного рядка.")
            return
        self.df_original.loc[idxs, "is_archived"] = True
        self._save_state()
        self.recalc_duplicate_marks(show_popup=False)
        self.apply_all_filters()

    def move_selected_from_archive(self):
        idxs = self.get_selected_indices()
        if not idxs:
            QMessageBox.information(self, "Архів", "Не вибрано жодного рядка.")
            return
        self.df_original.loc[idxs, "is_archived"] = False
        self._save_state()
        self.recalc_duplicate_marks(show_popup=False)
        self.apply_all_filters()

    def delete_selected_rows(self):
        idxs = self.get_selected_indices()
        if not idxs:
            QMessageBox.information(self, "Видалення", "Не вибрано жодного рядка.")
            return
        self.df_original.loc[idxs, "is_deleted"] = True
        self._save_state()
        self.recalc_duplicate_marks(show_popup=False)
        self.apply_all_filters()

    def restore_selected_rows(self):
        idxs = self.get_selected_indices()
        if not idxs:
            QMessageBox.information(self, "Відновлення", "Не вибрано жодного рядка.")
            return
        self.df_original.loc[idxs, "is_deleted"] = False
        self._save_state()
        self.recalc_duplicate_marks(show_popup=False)
        self.apply_all_filters()

    # -------------------- дублікати кнопкой --------------------

    def on_check_duplicates_clicked(self):
        old_count = len(self.duplicate_indices)
        self.recalc_duplicate_marks(show_popup=False)
        new_count = len(self.duplicate_indices)

        self.apply_all_filters()

        if new_count == 0:
            QMessageBox.information(self, "Дублікати", "Дублікати за ПІБ не виявлено.")
        else:
            msg = f"Виявлено {new_count} запис(ів)-дублікат(ів) за ПІБ.\n(Рядки підсвічені синім фоном.)"
            if new_count < old_count:
                msg += "\nЧастину дублікатів, ймовірно, було перенесено до 'Видалені'."
            QMessageBox.information(self, "Дублікати", msg)

    # -------------------- экспорт --------------------

    def _format_df_for_export(self, df: pd.DataFrame) -> pd.DataFrame:
        out = df.copy()
        for c in SERVICE_COLS:
            if c in out.columns:
                out = out.drop(columns=[c])
        for col in out.columns:
            if pd.api.types.is_datetime64_any_dtype(out[col]):
                out[col] = out[col].dt.strftime("%d.%m.%Y").fillna("")
            elif pd.api.types.is_bool_dtype(out[col]):
                out[col] = out[col].map({True: "Так", False: "Ні"})
        return out

    def export_file(self):
        if self.df_current is None or self.df_current.empty:
            QMessageBox.warning(self, "Експорт", "Немає даних для експорту.")
            return

        path, selected_filter = QFileDialog.getSaveFileName(
            self, "Зберегти результати фільтрації", "",
            "Word (*.docx);;Excel (*.xlsx);;CSV (*.csv)"
        )
        if not path:
            return

        try:
            df_out = self._format_df_for_export(self.df_current)

            if path.lower().endswith(".docx") or "Word" in selected_filter:
                doc = Document()
                section = doc.sections[0]
                section.orientation = WD_ORIENT.LANDSCAPE
                new_width, new_height = section.page_height, section.page_width
                section.page_width = new_width
                section.page_height = new_height

                table = doc.add_table(rows=1, cols=len(df_out.columns))
                table.style = "Table Grid"

                hdr_cells = table.rows[0].cells
                for j, col_name in enumerate(df_out.columns):
                    hdr_cells[j].text = str(col_name)

                for _, row in df_out.iterrows():
                    row_cells = table.add_row().cells
                    for j, col_name in enumerate(df_out.columns):
                        value = row[col_name]
                        row_cells[j].text = "" if pd.isna(value) else str(value)

                doc.save(path)

            elif path.lower().endswith(".xlsx") or "Excel" in selected_filter:
                df_out.to_excel(path, index=False)
            else:
                df_out.to_csv(path, index=False)

            QMessageBox.information(self, "Експорт", f"Файл збережено:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Помилка експорту", str(e))

    # -------------------- матчинг диалог --------------------

    def open_match_dialog(self):
        if self.df_original is None:
            QMessageBox.warning(self, "Аналіз", "Спочатку завантажте таблицю.")
            return
        dlg = MatchAnalysisDialog(parent=self, current_df=self.df_original)
        dlg.exec()


# ============================================================
#                  ТОЧКА ВХОДУ
# ============================================================

def main():
    app = QApplication(sys.argv)

    app.setStyleSheet("""
        QWidget {
            background-color: #f2f2f2;
            color: #111111;
            font-size: 14px;
        }
        QMainWindow {
            background-color: #f2f2f2;
        }
        QPushButton {
            background-color: #ffffff;
            border: 1px solid #bfbfbf;
            padding: 6px 10px;
            border-radius: 4px;
        }
        QPushButton:hover {
            background-color: #f5f5f5;
        }
        QPushButton:disabled {
            background-color: #e8e8e8;
            color: #999999;
        }
        QLineEdit, QComboBox {
            background-color: #ffffff;
            border: 1px solid #bfbfbf;
            border-radius: 4px;
            padding: 4px;
        }
        QListWidget {
            background-color: #ffffff;
            border: 1px solid #cccccc;
        }
        QTableView {
            background-color: #ffffff;
            alternate-background-color: #fafafa;
            gridline-color: #cccccc;
            selection-background-color: #cde7ff;
            selection-color: #000000;
        }
        QHeaderView::section {
            background-color: #e6e6e6;
            padding: 4px;
            border: 1px solid #c0c0c0;
        }
        QTabBar::tab {
            background: #e6e6e6;
            padding: 5px 12px;
            border: 1px solid #c0c0c0;
            border-bottom: none;
        }
        QTabBar::tab:selected {
            background: #ffffff;
        }
    """)

    window = MainWindow()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
